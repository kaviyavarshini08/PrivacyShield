from fastapi import FastAPI, HTTPException, status
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import os
import json
import logging
import fitz # PyMuPDF
from docx import Document as DocxDocument
import torch
from sentence_transformers import SentenceTransformer

from src.providers.manager import provider_manager
from src.ocr import perform_ocr
from src.cv_redactor import redact_image
from src.pdf_redactor import redact_pdf

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="PrivacyShield AI-Services Microservice", version="1.0.0")

# Embedding Model Global Initialization (loaded lazily on demand)
embedding_model = None

class EmbeddingsRequest(BaseModel):
    texts: List[str]

@app.post("/api/v1/embeddings")
async def generate_local_embeddings(req: EmbeddingsRequest):
    """
    Generates sentence-transformer embeddings using the loaded local model.
    """
    global embedding_model
    if embedding_model is None:
        try:
            logger.info("Embedding model not loaded. Attempting to load now...")
            enable_gpu = os.getenv("ENABLE_GPU_INFERENCE", "false").lower() == "true"
            device = "cuda" if (enable_gpu and torch.cuda.is_available()) else "cpu"
            embedding_model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2", device=device)
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Local embedding engine is unavailable: {str(e)}"
            )

    try:
        embeddings = embedding_model.encode(req.texts)
        result = [emb.tolist() for emb in embeddings]
        return {"embeddings": result}
    except Exception as e:
        logger.exception("Failed to generate local embeddings")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to generate embeddings: {str(e)}"
        )


class AnalyzeRequest(BaseModel):
    file_path: str
    content_type: str
    original_name: str

class RedactEntity(BaseModel):
    id: int
    entity_type: str
    text: str
    page_number: int
    bbox: List[float]
    start_char: int
    end_char: int

class RedactRequest(BaseModel):
    file_path: str
    content_type: str
    output_path: str
    entities: List[RedactEntity]

# Helper to extract text and bboxes from PDF using PyMuPDF (with OCR fallback for scanned PDFs)
def extract_pdf_data(pdf_path: str) -> List[Dict[str, Any]]:
    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        logger.error(f"Failed to open PDF: {e}")
        raise ValueError("Invalid PDF file format")
        
    if doc.is_encrypted:
        if not doc.authenticate(""):
            logger.warning(f"PDF {pdf_path} is password-protected.")
            doc.close()
            raise ValueError("Password-protected PDF. Please upload an unlocked PDF file.")

    pages_data = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        words = page.get_text("words") # yields list of: (x0, y0, x1, y1, "word", block_no, line_no, word_no)
        
        full_text = ""
        word_mappings = []
        for w in words:
            rect = [w[0], w[1], w[2], w[3]]
            word_text = w[4]
            
            start_idx = len(full_text)
            full_text += word_text + " "
            end_idx = len(full_text) - 1
            
            word_mappings.append({
                "word": word_text,
                "start": start_idx,
                "end": end_idx,
                "bbox": rect
            })
            
        # OCR scan for image-only, scanned, or image-embedded PDF pages
        page_images = page.get_images()
        if not full_text.strip() or len(page_images) > 0 or len(full_text.strip()) < 200:
            logger.info(f"Page {page_num + 1} has embedded images or sparse text. Running OCR scan...")
            try:
                dpi = 150
                pix = page.get_pixmap(dpi=dpi)
                temp_img_path = f"{pdf_path}_temp_page_{page_num}.png"
                pix.save(temp_img_path)
                ocr_text, ocr_mappings = perform_ocr(temp_img_path)
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)
                if ocr_text.strip() and (len(ocr_text.strip()) > len(full_text.strip()) or not full_text.strip()):
                    scale_x = page.rect.width / float(pix.width)
                    scale_y = page.rect.height / float(pix.height)
                    scaled_mappings = []
                    for m in ocr_mappings:
                        b = m["bbox"]
                        scaled_mappings.append({
                            "word": m["word"],
                            "start": m["start"],
                            "end": m["end"],
                            "bbox": [b[0] * scale_x, b[1] * scale_y, b[2] * scale_x, b[3] * scale_y]
                        })
                    full_text = ocr_text
                    word_mappings = scaled_mappings
            except Exception as ocr_err:
                logger.error(f"OCR fallback error on page {page_num + 1}: {ocr_err}")

        pages_data.append({
            "page_number": page_num + 1,
            "text": full_text,
            "word_mappings": word_mappings
        })


    doc.close()
    return pages_data

def find_bbox_for_offset(start: int, end: int, word_mappings: List[Dict]) -> List[float]:
    bboxes = []
    for m in word_mappings:
        # Check overlap
        if max(start, m["start"]) <= min(end, m["end"]):
            bboxes.append(m["bbox"])
            
    if not bboxes:
        return [0.0, 0.0, 0.0, 0.0]
        
    # Take union of overlapped word bboxes
    x0 = min(b[0] for b in bboxes)
    y0 = min(b[1] for b in bboxes)
    x1 = max(b[2] for b in bboxes)
    y1 = max(b[3] for b in bboxes)
    return [x0, y0, x1, y1]

@app.post("/api/v1/analyze")
async def analyze_document(req: AnalyzeRequest):
    """
    Parses and scans a document for PII using Presidio + OCR + PDF text position extraction.
    """
    file_path = req.file_path
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found on storage mount.")
        
    content_type = req.content_type.lower()
    file_ext = os.path.splitext(req.original_name)[1].lower()
    
    entities = []
    
    try:
        # 1. Image OCR flow
        if "image" in content_type or file_ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
            text, mappings = perform_ocr(file_path)
            raw_entities = provider_manager.get_provider().analyze(text)
            
            for ent in raw_entities:
                # Resolve image bounding box coordinates
                bbox = find_bbox_for_offset(ent["start_char"], ent["end_char"], mappings)
                ent["bbox"] = bbox
                ent["page_number"] = 1
                entities.append(ent)
                
        # 2. PDF text positioning flow
        elif content_type == "application/pdf" or file_ext == ".pdf":
            pages_data = extract_pdf_data(file_path)
            for page in pages_data:
                raw_entities = provider_manager.get_provider().analyze(page["text"])
                for ent in raw_entities:
                    bbox = find_bbox_for_offset(ent["start_char"], ent["end_char"], page["word_mappings"])
                    ent["bbox"] = bbox
                    ent["page_number"] = page["page_number"]
                    entities.append(ent)
                    
        # 3. DOCX text flow
        elif file_ext == ".docx":
            doc = DocxDocument(file_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            raw_entities = provider_manager.get_provider().analyze(full_text)
            for ent in raw_entities:
                ent["bbox"] = [0.0, 0.0, 0.0, 0.0]
                ent["page_number"] = 1
                entities.append(ent)
                
        # 4. Plain Text / CSV flow
        else:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            raw_entities = provider_manager.get_provider().analyze(text)
            for ent in raw_entities:
                ent["bbox"] = [0.0, 0.0, 0.0, 0.0]
                ent["page_number"] = 1
                entities.append(ent)
                
        return {"entities": entities}
        
    except Exception as e:
        logger.exception("Failed to analyze document")
        raise HTTPException(
            status_code=500,
            detail=f"Analysis pipeline error: {str(e)}"
        )

@app.post("/api/v1/redact")
async def redact_document(req: RedactRequest):
    """
    Executes redactions on the target file and outputs the sanitized document.
    """
    if not os.path.exists(req.file_path):
        raise HTTPException(status_code=404, detail="Original file not found")
        
    content_type = req.content_type.lower()
    file_ext = os.path.splitext(req.file_path)[1].lower()
    
    entities_data = [e.model_dump() for e in req.entities]
    
    try:
        # 1. PDF Redaction
        if content_type == "application/pdf" or file_ext == ".pdf":
            redact_pdf(req.file_path, req.output_path, entities_data)
            
        # 2. Image Redaction
        elif "image" in content_type or file_ext in [".png", ".jpg", ".jpeg"]:
            bboxes = [ent["bbox"] for ent in entities_data]
            redact_image(req.file_path, req.output_path, bboxes, mode="blur")
            
        # 3. DOCX Redaction
        elif file_ext == ".docx":
            doc = DocxDocument(req.file_path)
            redacted_words = [ent["text"] for ent in entities_data]
            
            for p in doc.paragraphs:
                for word in redacted_words:
                    if word in p.text:
                        p.text = p.text.replace(word, "[REDACTED]")
            doc.save(req.output_path)
            
        # 4. Text Redaction (TXT/CSV)
        else:
            with open(req.file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
                
            redacted_words = [ent["text"] for ent in entities_data]
            for word in redacted_words:
                text = text.replace(word, "[REDACTED]")
                
            with open(req.output_path, "w", encoding="utf-8") as f:
                f.write(text)
                
        return {"status": "success", "redacted_path": req.output_path}
        
    except Exception as e:
        logger.exception("Failed to apply redactions")
        raise HTTPException(
            status_code=500,
            detail=f"Redaction failed: {str(e)}"
        )

@app.get("/health")
def health():
    return {"status": "healthy"}
